from fastapi import APIRouter, HTTPException, Response
from fastapi.responses import FileResponse
from docx import Document
from io import BytesIO
from src.models import Cv
from sqlalchemy.orm import Session
from fastapi import Depends
from src.database import get_db

router = APIRouter()

logging.basicConfig(level=logging.INFO)
logging.info("Cargando rutas de cv...")


@router.post("/generate-docx/")
async def generate_docx(cv_id: str, db: Session = Depends(get_db)):
    """
    Endpoint para generar un CV en formato Word (.docx) a partir de un ID de CV.
    """
    # Buscar el CV en la base de datos por su ID
    cv = db.query(Cv).filter(Cv.id == cv_id).first()
    
    if not cv:
        raise HTTPException(status_code=404, detail="CV no encontrado")
    
    # Crear un nuevo documento Word
    document = Document()
    
    # Agregar título del CV
    document.add_heading("Currículum Vitae", level=1)
    
    # Agregar sección de datos básicos
    document.add_heading("Datos Personales", level=2)
    document.add_paragraph(f"Nombre: {cv.user.name}")  # Asumiendo que el modelo User tiene un campo "name"
    document.add_paragraph(f"Teléfono: {cv.telephone}")
    document.add_paragraph(f"Rol: {cv.role}")
    
    # Agregar resumen
    document.add_heading("Resumen", level=2)
    document.add_paragraph(cv.summary)
    
    # Agregar experiencia laboral
    document.add_heading("Experiencia Laboral", level=2)
    for exp in cv.experiences:
        document.add_paragraph(f"{exp.role} - {exp.company}", style="List Bullet")
        document.add_paragraph(f"Descripción: {exp.description}")
        document.add_paragraph(f"Fecha: {exp.start_date.strftime('%Y-%m')} - {exp.end_date.strftime('%Y-%m')}")
    
    # Agregar habilidades técnicas
    document.add_heading("Habilidades Técnicas", level=2)
    for skill in cv.skills:
        document.add_paragraph(f"{skill.title} - Nivel: {skill.level}", style="List Bullet")
    
    # Agregar habilidades blandas
    document.add_heading("Habilidades Blandas", level=2)
    for soft_skill in cv.soft_skills:
        document.add_paragraph(soft_skill.title, style="List Bullet")
    
    # Guardar el documento en memoria
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    # Devolver el archivo como respuesta descargable
    headers = {
        "Content-Disposition": "attachment; filename=cv.docx"
    }
    return Response(content=file_stream.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)


